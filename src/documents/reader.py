"""Safe, reusable text extraction for PDF, DOCX, and TXT documents."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO, Union

import fitz
from docx import Document

DocumentSource = Union[str, Path, bytes, bytearray, BinaryIO]
SUPPORTED_EXTENSIONS = {".pdf": "pdf", ".docx": "docx", ".txt": "txt"}


class DocumentReadError(ValueError):
    """Raised when a supported document cannot be read or yields no text."""


class UnsupportedDocumentTypeError(DocumentReadError):
    """Raised when the document type is not PDF, DOCX, or TXT."""


def _source_name(source: DocumentSource, filename: str | None = None) -> str:
    if filename:
        return filename
    if isinstance(source, (str, Path)):
        return str(source)
    return str(getattr(source, "name", ""))


def detect_document_type(
    source: DocumentSource, filename: str | None = None
) -> str:
    """Detect a supported type from a filename and, when needed, file bytes."""
    name = _source_name(source, filename)
    extension = Path(name).suffix.lower()
    if extension in SUPPORTED_EXTENSIONS:
        return SUPPORTED_EXTENSIONS[extension]

    data = _read_bytes(source)
    if data.startswith(b"%PDF-"):
        return "pdf"
    if data.startswith(b"PK\x03\x04"):
        return "docx"

    if extension:
        raise UnsupportedDocumentTypeError(
            f"Unsupported document type '{extension}'. Use PDF, DOCX, or TXT."
        )
    raise UnsupportedDocumentTypeError(
        "Could not determine the document type. Provide a .pdf, .docx, or .txt filename."
    )


def _read_bytes(source: DocumentSource) -> bytes:
    if isinstance(source, (str, Path)):
        path = Path(source)
        if not path.is_file():
            raise DocumentReadError(f"Document not found: {path}")
        try:
            return path.read_bytes()
        except OSError as exc:
            raise DocumentReadError(f"Could not read document '{path}': {exc}") from exc
    if isinstance(source, (bytes, bytearray)):
        return bytes(source)
    if hasattr(source, "read"):
        try:
            position = source.tell() if hasattr(source, "tell") else None
            data = source.read()
            if position is not None and hasattr(source, "seek"):
                source.seek(position)
            if isinstance(data, str):
                return data.encode("utf-8")
            return bytes(data)
        except (OSError, TypeError, ValueError) as exc:
            raise DocumentReadError(f"Could not read the uploaded document: {exc}") from exc
    raise DocumentReadError("Document source must be a path, bytes, or a readable stream.")


def extract_pdf_text(data: bytes) -> str:
    """Extract text from all pages of a PDF byte sequence."""
    try:
        with fitz.open(stream=data, filetype="pdf") as document:
            return "\n\n".join(page.get_text("text") for page in document)
    except Exception as exc:
        raise DocumentReadError(f"PDF extraction failed: {exc}") from exc


def extract_docx_text(data: bytes) -> str:
    """Extract paragraphs and table-cell text from a DOCX byte sequence."""
    try:
        document = Document(io.BytesIO(data))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(blocks)
    except Exception as exc:
        raise DocumentReadError(f"DOCX extraction failed: {exc}") from exc


def extract_txt_text(data: bytes) -> str:
    """Decode plain text using common encodings, failing clearly if none work."""
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError("TXT decoding failed: unsupported text encoding.")


def normalize_text(text: str) -> str:
    """Normalize whitespace while retaining paragraph boundaries."""
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    lines = [re.sub(r"[^\S\n]+", " ", line).strip() for line in text.split("\n")]
    normalized = "\n".join(lines)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def read_document(
    source: DocumentSource, filename: str | None = None
) -> str:
    """Read one document and return normalized text.

    Errors are raised per document so callers processing a batch can catch the
    error, record it, and continue with remaining resumes.
    """
    data = _read_bytes(source)
    if not data:
        raise DocumentReadError("The document is empty.")
    document_type = detect_document_type(source, filename)
    extractors = {
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "txt": extract_txt_text,
    }
    text = normalize_text(extractors[document_type](data))
    if not text:
        raise DocumentReadError(
            "No readable text was extracted; the document may be blank or image-only."
        )
    return text
